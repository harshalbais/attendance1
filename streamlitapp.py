import streamlit as st
import pandas as pd
from datetime import datetime
from docx import Document
from io import BytesIO

# --- Page Config ---
st.set_page_config(page_title="Class Attendance", layout="centered")

# --- Imaginary Student List ---
student_list = [
    {"Name": "Aarav Singh", "Roll Number": "101"},
    {"Name": "Diya Sharma", "Roll Number": "102"},
    {"Name": "Rohan Mehta", "Roll Number": "103"},
    {"Name": "Sanya Verma", "Roll Number": "104"},
    {"Name": "Yash Patil", "Roll Number": "105"},
]

# --- Session State ---
if "present_list" not in st.session_state:
    st.session_state.present_list = []

if "absent_list" not in st.session_state:
    st.session_state.absent_list = []

st.title("📋 Class Attendance System")

st.subheader("👩‍🏫 Mark Attendance")

# --- Attendance Form ---
with st.form("attendance_form"):
    marked_attendance = {}
    for student in student_list:
        key = f"{student['Roll Number']}_status"
        status = st.checkbox(
            f"{student['Roll Number']} - {student['Name']}",
            key=key,
            value=True
        )
        marked_attendance[student["Roll Number"]] = "Present" if status else "Absent"

    submitted = st.form_submit_button("✅ Submit Attendance")

    if submitted:
        today = datetime.now().strftime("%Y-%m-%d")
        st.session_state.present_list.clear()
        st.session_state.absent_list.clear()

        for student in student_list:
            record = {
                "Date": today,
                "Name": student["Name"],
                "Roll Number": student["Roll Number"]
            }
            if marked_attendance[student["Roll Number"]] == "Present":
                st.session_state.present_list.append(record)
            else:
                st.session_state.absent_list.append(record)

        st.success("🎉 Attendance submitted successfully!")

# --- Display Tables ---
def show_table(title, data_list, color):
    if data_list:
        st.subheader(f"{title}")
        df = pd.DataFrame(data_list)
        st.dataframe(df, use_container_width=True)
        return df
    return None

df_present = show_table("✅ Present Students", st.session_state.present_list, "green")
df_absent = show_table("❌ Absent Students", st.session_state.absent_list, "red")

# --- Create Word file ---
def create_word_file(present_df, absent_df):
    doc = Document()
    doc.add_heading("Class Attendance Report", 0)
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph("")

    if not present_df.empty:
        doc.add_heading("✅ Present Students", level=1)
        table = doc.add_table(rows=1, cols=len(present_df.columns))
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(present_df.columns):
            hdr_cells[i].text = col
        for _, row in present_df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
        doc.add_paragraph("")

    if not absent_df.empty:
        doc.add_heading("❌ Absent Students", level=1)
        table = doc.add_table(rows=1, cols=len(absent_df.columns))
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(absent_df.columns):
            hdr_cells[i].text = col
        for _, row in absent_df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Download Word Button ---
if submitted:
    word_file = create_word_file(df_present if df_present is not None else pd.DataFrame(),
                                  df_absent if df_absent is not None else pd.DataFrame())

    st.download_button(
        label="📥 Download Attendance as Word File (.docx)",
        data=word_file,
        file_name="attendance_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
